"""Covers pipeline/word/merge.py - the DOCX merge/insert feature
(01.09.2026, Michael: "Jetzt noch das ganze für *.docx. Kann man
überhaupt mittlerweile ca. 2000 docx zusammenführen?"). Mirrors
tests/test_pdf_merge.py's coverage where the same concepts apply
(order, dedup counting, cancellation keeps a partial result, hard vs.
soft per-file failure) and adds DOCX-specific coverage for the
two-level batching architecture Michael explicitly confirmed
("automatisch batchen").

Fixtures used here are the plain synthetic .docx files built inline
below via python-docx (see _make_docx()) - the module-level "only
pipeline/word/merge.py imports docx" rule (see that module's docstring)
has the same narrow, documented test-fixture exception
tests/test_pdf_ico_mode.py already established for fitz/PyMuPDF.
"""
from __future__ import annotations

from pathlib import Path

import docx
import docx.oxml.ns
import pytest
from docx.enum.section import WD_SECTION

from pipeline.word.merge import DEFAULT_BATCH_SIZE, WordMergeStats, merge_docx_files


def _make_docx(path: Path, text: str) -> Path:
    document = docx.Document()
    document.add_paragraph(text)
    document.save(str(path))
    return path


def _paragraph_texts(path: Path) -> list[str]:
    return [p.text for p in docx.Document(str(path)).paragraphs]


def test_merge_preserves_source_order(tmp_path: Path) -> None:
    a = _make_docx(tmp_path / "a.docx", "First document")
    b = _make_docx(tmp_path / "b.docx", "Second document")
    c = _make_docx(tmp_path / "c.docx", "Third document")
    destination = tmp_path / "out.docx"

    stats = merge_docx_files([a, b, c], destination)

    texts = [t for t in _paragraph_texts(destination) if t]
    assert texts == ["First document", "Second document", "Third document"]
    assert stats == WordMergeStats(segments=3, files_processed=3, batches=0, cancelled=False, warnings=[])


def test_merge_inserts_a_new_page_section_break_between_sources(tmp_path: Path) -> None:
    """Since 03.09.2026 a SECTION break (new page), not a plain page break -
    see _append_as_own_section(): headers/footers are section properties,
    so every source needs its own section to keep them. No explicit
    <w:br w:type="page"/> any more either, or there would be an empty page
    between the two."""
    a = _make_docx(tmp_path / "a.docx", "First document")
    b = _make_docx(tmp_path / "b.docx", "Second document")
    destination = tmp_path / "out.docx"

    merge_docx_files([a, b], destination)

    document = docx.Document(str(destination))
    assert len(document.sections) == 2
    assert all(section.start_type == WD_SECTION.NEW_PAGE for section in document.sections)
    assert 'w:type="page"' not in document.element.body.xml


def test_single_source_has_a_single_section(tmp_path: Path) -> None:
    a = _make_docx(tmp_path / "a.docx", "Only document")
    destination = tmp_path / "out.docx"

    merge_docx_files([a], destination)

    document = docx.Document(str(destination))
    assert len(document.sections) == 1
    assert [p.text for p in document.paragraphs if p.text] == ["Only document"]


def _make_docx_with_header(
    path: Path, text: str, header: str, footer: str | None = None, second_section_header: str | None = None,
) -> Path:
    document = docx.Document()
    document.add_paragraph(text)
    document.sections[0].header.paragraphs[0].text = header
    if footer is not None:
        document.sections[0].footer.paragraphs[0].text = footer
    if second_section_header is not None:
        section = document.add_section(WD_SECTION.NEW_PAGE)
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = second_section_header
        document.add_paragraph(text + " (part 2)")
    document.save(str(path))
    return path


def _section_headers_and_footers(path: Path) -> list[tuple[str, str]]:
    document = docx.Document(str(path))
    return [(s.header.paragraphs[0].text, s.footer.paragraphs[0].text) for s in document.sections]


def test_every_source_keeps_its_own_header_and_footer(tmp_path: Path) -> None:
    """Regression guard for 03.09.2026 (Michael: "Beim Zusammenführen der
    Worddokumente wird [...] der Header des ersten Dokuments übernommen.
    Auf jeden Fall ist dort auf jeder Seite der gleiche Header."): plain
    docxcompose drops every appended document's section properties, so
    the first file's header/footer ended up on every page."""
    a = _make_docx_with_header(tmp_path / "a.docx", "Doc A", "Header A", "Footer A")
    b = _make_docx_with_header(tmp_path / "b.docx", "Doc B", "Header B", "Footer B")
    c = _make_docx_with_header(tmp_path / "c.docx", "Doc C", "Header C")
    destination = tmp_path / "out.docx"

    merge_docx_files([a, b, c], destination)

    # Doc C has no footer of its own: it must get a BLANK one, not inherit
    # Footer B from the section before it (OOXML's default for a section
    # without a footer reference).
    assert _section_headers_and_footers(destination) == [
        ("Header A", "Footer A"), ("Header B", "Footer B"), ("Header C", ""),
    ]


def test_multi_section_source_keeps_all_of_its_headers(tmp_path: Path) -> None:
    a = _make_docx_with_header(tmp_path / "a.docx", "Doc A", "Header A")
    b = _make_docx_with_header(tmp_path / "b.docx", "Doc B", "Header B", second_section_header="Header B / 2")
    destination = tmp_path / "out.docx"

    merge_docx_files([a, b], destination)

    assert [h for h, _ in _section_headers_and_footers(destination)] == ["Header A", "Header B", "Header B / 2"]
    texts = [t for t in _paragraph_texts(destination) if t]
    assert texts == ["Doc A", "Doc B", "Doc B (part 2)"]


def test_header_images_are_copied_with_the_header(tmp_path: Path) -> None:
    from PIL import Image

    logo = tmp_path / "logo.png"
    Image.new("RGB", (40, 20), "red").save(logo)
    a = _make_docx_with_header(tmp_path / "a.docx", "Doc A", "Header A")
    b = docx.Document()
    b.add_paragraph("Doc B")
    b.sections[0].header.paragraphs[0].add_run("Header B ").add_picture(str(logo))
    b.save(str(tmp_path / "b.docx"))
    destination = tmp_path / "out.docx"

    merge_docx_files([a, tmp_path / "b.docx"], destination)

    merged = docx.Document(str(destination))
    header = merged.sections[1].header
    blips = header._element.xpath(".//a:blip")
    assert len(blips) == 1
    rid = blips[0].get(docx.oxml.ns.qn("r:embed"))
    assert rid in header.part.rels  # the image relationship was carried over, not left dangling


def test_headers_survive_two_level_batching(tmp_path: Path) -> None:
    sources = [_make_docx_with_header(tmp_path / f"d{i}.docx", f"Doc {i}", f"Header {i}") for i in range(5)]
    destination = tmp_path / "out.docx"

    stats = merge_docx_files(sources, destination, batch_size=2)

    assert stats.batches == 3
    assert [h for h, _ in _section_headers_and_footers(destination)] == [f"Header {i}" for i in range(5)]


def test_same_file_listed_twice_counts_once_in_files_processed(tmp_path: Path) -> None:
    a = _make_docx(tmp_path / "a.docx", "Repeated document")
    destination = tmp_path / "out.docx"

    stats = merge_docx_files([a, a], destination)

    assert stats.segments == 2  # appended twice - the "insert same file again" use case
    assert stats.files_processed == 1  # but it's the same resolved path


def test_empty_sources_list_raises(tmp_path: Path) -> None:
    with pytest.raises(ValueError):
        merge_docx_files([], tmp_path / "out.docx")


def test_destination_parent_directory_is_created(tmp_path: Path) -> None:
    a = _make_docx(tmp_path / "a.docx", "Doc")
    destination = tmp_path / "nested" / "deeper" / "out.docx"

    merge_docx_files([a], destination)

    assert destination.exists()


def test_unopenable_source_aborts_the_whole_merge(tmp_path: Path) -> None:
    a = _make_docx(tmp_path / "a.docx", "Good document")
    bad = tmp_path / "bad.docx"
    bad.write_text("not a docx at all")
    destination = tmp_path / "out.docx"

    with pytest.raises(ValueError, match="bad.docx"):
        merge_docx_files([a, bad], destination)
    assert not destination.exists()


def test_append_failure_is_skipped_with_a_warning_not_aborted(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    a = _make_docx(tmp_path / "a.docx", "First document")
    b = _make_docx(tmp_path / "b.docx", "Problem document")
    c = _make_docx(tmp_path / "c.docx", "Third document")
    destination = tmp_path / "out.docx"

    from docxcompose.composer import Composer

    original_append = Composer.append
    call_count = {"n": 0}

    def flaky_append(self, doc, remove_property_fields=True):
        call_count["n"] += 1
        if call_count["n"] == 1:  # the second source (b.docx) - first call to append()
            raise RuntimeError("simulated docxcompose SmartArt failure")
        return original_append(self, doc, remove_property_fields=remove_property_fields)

    monkeypatch.setattr(Composer, "append", flaky_append)

    stats = merge_docx_files([a, b, c], destination)

    texts = [t for t in _paragraph_texts(destination) if t]
    assert texts == ["First document", "Third document"]  # b.docx skipped, merge continued
    assert stats.segments == 2
    assert len(stats.warnings) == 1
    assert "b.docx" in stats.warnings[0]


def test_below_batch_size_uses_a_single_pass_with_no_batches(tmp_path: Path) -> None:
    sources = [_make_docx(tmp_path / f"f{i}.docx", f"Doc {i}") for i in range(3)]
    destination = tmp_path / "out.docx"

    stats = merge_docx_files(sources, destination, batch_size=10)

    assert stats.batches == 0


def test_above_batch_size_uses_two_level_batching_and_preserves_order(tmp_path: Path) -> None:
    sources = [_make_docx(tmp_path / f"f{i}.docx", f"Doc {i}") for i in range(7)]
    destination = tmp_path / "out.docx"

    stats = merge_docx_files(sources, destination, batch_size=3)

    assert stats.batches == 3  # ceil(7/3)
    assert stats.segments == 7
    texts = [t for t in _paragraph_texts(destination) if t]
    assert texts == [f"Doc {i}" for i in range(7)]


def test_default_batch_size_is_used_when_not_specified(tmp_path: Path) -> None:
    sources = [_make_docx(tmp_path / f"f{i}.docx", f"Doc {i}") for i in range(3)]
    destination = tmp_path / "out.docx"

    stats = merge_docx_files(sources, destination)

    assert DEFAULT_BATCH_SIZE > 3
    assert stats.batches == 0  # comfortably below the default batch size


def test_cancellation_within_the_first_batch_keeps_only_what_completed(tmp_path: Path) -> None:
    sources = [_make_docx(tmp_path / f"f{i}.docx", f"Doc {i}") for i in range(6)]
    destination = tmp_path / "out.docx"

    calls = {"n": 0}

    def should_cancel() -> bool:
        calls["n"] += 1
        return calls["n"] > 2  # stop after the first 2 files of batch 1

    stats = merge_docx_files(sources, destination, batch_size=3, should_cancel=should_cancel)

    assert stats.cancelled is True
    texts = [t for t in _paragraph_texts(destination) if t]
    assert texts == ["Doc 0", "Doc 1"]


def test_cancellation_after_the_first_batch_keeps_every_completed_batch(tmp_path: Path) -> None:
    """Regression guard: an earlier version of this code only kept the
    FIRST completed batch on cancellation (the second-level "merge the
    chunks" pass polled the same should_cancel, which - being a
    threading.Event in real use - stays set once tripped, so it
    immediately bailed out and silently discarded every batch after the
    first). Cancelling partway through batch 2 of 3 must still keep both
    fully-completed batch 1 AND the partial work done in batch 2, not
    just batch 1.
    """
    sources = [_make_docx(tmp_path / f"f{i}.docx", f"Doc {i}") for i in range(9)]
    destination = tmp_path / "out.docx"

    calls = {"n": 0}

    def should_cancel() -> bool:
        calls["n"] += 1
        return calls["n"] > 4  # completes batch 1 (3 files) + 1 file into batch 2

    stats = merge_docx_files(sources, destination, batch_size=3, should_cancel=should_cancel)

    assert stats.cancelled is True
    assert stats.batches == 2
    texts = [t for t in _paragraph_texts(destination) if t]
    assert texts == ["Doc 0", "Doc 1", "Doc 2", "Doc 3"]


def test_cancellation_before_any_file_is_appended_raises(tmp_path: Path) -> None:
    sources = [_make_docx(tmp_path / f"f{i}.docx", f"Doc {i}") for i in range(3)]
    destination = tmp_path / "out.docx"

    with pytest.raises(ValueError):
        merge_docx_files(sources, destination, should_cancel=lambda: True)


def test_progress_callback_reports_file_and_batch_messages(tmp_path: Path) -> None:
    sources = [_make_docx(tmp_path / f"f{i}.docx", f"Doc {i}") for i in range(5)]
    destination = tmp_path / "out.docx"
    messages: list[str] = []

    merge_docx_files(sources, destination, batch_size=2, progress_callback=messages.append)

    assert any("Batch" in m for m in messages)
    assert any("f0.docx" in m for m in messages)
    assert any("Zwischenergebnisse" in m for m in messages)


def test_partial_append_failure_leaves_nothing_of_the_skipped_document_behind(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    """03.09.2026 (Michael: "Das sind solche Fehler die nicht so gut
    auffallen."): docxcompose inserts paragraph by paragraph and can fail
    halfway through a document. The warning says "übersprungen", so the
    result must not contain the first half of that document either - the
    body is rolled back to its exact state before the failed append."""
    a = _make_docx_with_header(tmp_path / "a.docx", "Doc A", "Header A")
    b = _make_docx_with_header(tmp_path / "b.docx", "Half of B", "Header B")
    c = _make_docx_with_header(tmp_path / "c.docx", "Doc C", "Header C")
    destination = tmp_path / "out.docx"

    from docxcompose.composer import Composer

    original_append = Composer.append
    calls = {"n": 0}

    def half_then_fail(self, doc, remove_property_fields=True):
        calls["n"] += 1
        original_append(self, doc, remove_property_fields=remove_property_fields)  # content IS inserted ...
        if calls["n"] == 1:
            raise RuntimeError("simulated failure after the content was already inserted")

    monkeypatch.setattr(Composer, "append", half_then_fail)

    stats = merge_docx_files([a, b, c], destination)

    texts = [t for t in _paragraph_texts(destination) if t]
    assert texts == ["Doc A", "Doc C"]
    assert stats.segments == 2
    assert len(stats.warnings) == 1 and "b.docx" in stats.warnings[0]
    assert [h for h, _ in _section_headers_and_footers(destination)] == ["Header A", "Header C"]


def _make_docx_with_header_and_body_image(path: Path, text: str, header_color, body_color) -> Path:
    """Header AND body each carry a distinct, unique image - see
    test_merging_many_documents_with_distinct_header_and_body_images_produces_no_duplicate_media_parts()."""
    from PIL import Image

    header_img = path.with_name(path.stem + "_header.png")
    body_img = path.with_name(path.stem + "_body.png")
    Image.new("RGB", (10, 10), header_color).save(header_img)
    Image.new("RGB", (12, 12), body_color).save(body_img)

    document = docx.Document()
    run = document.add_paragraph().add_run()
    run.add_picture(str(body_img))
    document.add_paragraph(text)
    document.sections[0].header.paragraphs[0].add_run(f"Header for {text} ").add_picture(str(header_img))
    document.save(str(path))
    return path


def test_merging_many_documents_with_distinct_header_and_body_images_produces_no_duplicate_media_parts(
    tmp_path: Path,
) -> None:
    """Regression guard for 03.09.2026 (Michael, real merged file: "ist
    defekt und kann deshalb nicht geöffnet werden [...] Soll LibreOffice
    die Datei reparieren?"): copying a header/footer image via
    Composer.add_relationship()'s generic, package-wide filename scan
    picks numbers independently from python-docx's own
    Package.image_parts registry (used for BODY images by docxcompose
    itself) - the two can hand out the SAME "next free" image partname,
    producing a zip with two entries of the same name (only visible with
    per-document-UNIQUE images; identical images across documents get
    deduplicated by sha1 before the collision would ever occur, which is
    why this needs its own fixture rather than reusing the shared-logo
    tests above).
    """
    import zipfile

    sources = [
        _make_docx_with_header_and_body_image(
            tmp_path / f"d{i}.docx", f"Doc {i}", (i * 20 % 255, 0, 0), (0, i * 20 % 255, 0),
        )
        for i in range(8)
    ]
    destination = tmp_path / "out.docx"

    merge_docx_files(sources, destination)

    with zipfile.ZipFile(destination) as archive:
        names = archive.namelist()
    assert len(names) == len(set(names)), f"duplicate zip entries: {[n for n in set(names) if names.count(n) > 1]}"

    # And every merged section still shows its own, distinct header.
    merged = docx.Document(str(destination))
    assert [s.header.paragraphs[0].text for s in merged.sections] == [f"Header for Doc {i} " for i in range(8)]
