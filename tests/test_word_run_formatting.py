"""Covers three related 03.09.2026 regressions Michael reported from real
translated documents (side-by-side mode, but the first two affect normal
mode identically - see docx_engine.py's docstrings for the full story):

"Ich sehe noch Unterschiede bei den Fonts, Links werden bei der
Übersetzung scheinbar kaputt gemacht und wie im Bild sieht man das ein
Aufzählungszeichen scheinbar übersetzt wurde... Das ist bei diesem
Aufzählungszeichen ein Muster und kommt immer wieder vor."

1. A translated run's original font/size (<w:rFonts>/<w:sz>, not just
   bold/italic/underline) must survive - previously every rebuilt run got
   a from-scratch <w:rPr> that discarded everything except b/i/u.
2. A translated hyperlink must keep its visual styling (<w:rStyle
   w:val="Hyperlink"/> plus its own font/size) - previously
   _build_hyperlink_element() wrote no <w:rPr> at all.
3. A <w:sym/> (symbol-font bullet, e.g. Wingdings) sitting in its own run
   must survive translation instead of silently vanishing - previously
   _walk_run() didn't recognize <w:sym/> at all, so it was never even
   represented as a WordRun.

Builds real .docx fixtures via python-docx (+ raw oxml for the hyperlink/
symbol elements python-docx has no high-level API for), runs them through
the real DocxEngine + translate_document() pipeline with a fake provider,
then reads the result back with raw lxml (python-docx has no rPr-level
API either) to assert on the actual rebuilt XML - not just that no
exception was raised.
"""
from __future__ import annotations

import zipfile
from pathlib import Path

import docx
import pytest
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

from pipeline.translation.base import TranslationResult
from pipeline.word.docx_engine import DocxEngine, _w
from pipeline.word.translate_document import translate_document

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class _FakeProvider:
    """Prefixes the (plain-text/tagged) HTML so translated text is
    trivially distinguishable - same convention as
    tests/test_word_side_by_side.py's _FakeProvider."""

    def translate_html(self, html: str, target_lang: str, source_lang: str | None = None) -> TranslationResult:
        return TranslationResult(text=f"[{target_lang}] {html}", source_lang=source_lang or "en", target_lang=target_lang, provider="fake")


def _add_hyperlink(paragraph, text: str, url: str, font_name: str | None = None, size_half_pt: int | None = None):
    """Build a real <w:hyperlink> the way Word does - rStyle="Hyperlink"
    plus (for this test) an explicit font/size, exactly like the run
    formatting Michael's real documents had on their hyperlinks -
    python-docx itself has no API for this."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    if font_name:
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rpr.append(rfonts)
    if size_half_pt:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size_half_pt))
        rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return url


def _add_symbol_run(paragraph, font: str, char_code: str) -> None:
    """<w:r><w:sym w:font="..." w:char="..."/></w:r> - a symbol-font
    character reference (e.g. a Wingdings bullet), structurally distinct
    from <w:t> text. python-docx has no API for this either."""
    run = OxmlElement("w:r")
    sym = OxmlElement("w:sym")
    sym.set(qn("w:font"), font)
    sym.set(qn("w:char"), char_code)
    run.append(sym)
    paragraph._p.append(run)


def test_font_and_size_survive_translation(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = docx.Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Please read this carefully.")
    run.font.name = "Montserrat"
    run.font.size = Pt(12)
    document.save(str(source))

    destination = tmp_path / "result.docx"
    engine = DocxEngine()
    engine.open(str(source))
    translate_document(engine, _FakeProvider(), [], "DE", "en")
    engine.save(str(destination))

    result = docx.Document(str(destination))
    assert result.paragraphs[0].text == "[DE] Please read this carefully."

    rpr = result.paragraphs[0].runs[0]._r.find(_w("rPr"))
    assert rpr is not None, "translated run lost its <w:rPr> entirely"
    rfonts = rpr.find(_w("rFonts"))
    assert rfonts is not None, "translated run lost <w:rFonts> - the font-loss regression"
    assert rfonts.get(_w("ascii")) == "Montserrat"
    sz = rpr.find(_w("sz"))
    assert sz is not None
    assert sz.get(_w("val")) == "24"  # 12pt * 2 (half-points)


def test_hyperlink_keeps_its_visual_style_after_translation(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = docx.Document()
    paragraph = document.add_paragraph("See ")
    url = _add_hyperlink(paragraph, "our site", "https://example.com", font_name="Montserrat ExtraBold", size_half_pt=24)
    paragraph.add_run(" for details.")
    document.save(str(source))

    destination = tmp_path / "result.docx"
    engine = DocxEngine()
    engine.open(str(source))
    translate_document(engine, _FakeProvider(), [], "DE", "en")
    engine.save(str(destination))

    result_bytes = Path(destination).read_bytes()
    with zipfile.ZipFile(destination) as archive:
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)

    hyperlinks = root.findall(f".//{{{_W_NS}}}hyperlink")
    assert len(hyperlinks) == 1
    inner_run = hyperlinks[0].find(f"{{{_W_NS}}}r")
    rpr = inner_run.find(f"{{{_W_NS}}}rPr")
    assert rpr is not None, "translated hyperlink lost its <w:rPr> entirely - renders as plain black text"

    rstyle = rpr.find(f"{{{_W_NS}}}rStyle")
    assert rstyle is not None
    assert rstyle.get(f"{{{_W_NS}}}val") == "Hyperlink"

    rfonts = rpr.find(f"{{{_W_NS}}}rFonts")
    assert rfonts is not None
    assert rfonts.get(f"{{{_W_NS}}}ascii") == "Montserrat ExtraBold"

    sz = rpr.find(f"{{{_W_NS}}}sz")
    assert sz is not None
    assert sz.get(f"{{{_W_NS}}}val") == "24"

    # The fake provider prefixes the whole HTML string with "[DE] " -
    # landing on the plain-text run BEFORE the <a> tag, not inside it, so
    # the hyperlink's own display text is untouched here. That's expected
    # HTML-structure-preserving behavior, not part of what this test is
    # checking (the rPr/style assertions above are the actual point).
    text = "".join(t.text or "" for t in inner_run.findall(f"{{{_W_NS}}}t"))
    assert text == "our site"


def test_symbol_bullet_survives_translation(tmp_path: Path) -> None:
    """03.09.2026 (Michael, explicitly flagged as a recurring pattern):
    a <w:sym/> bullet in its own run, immediately followed by translatable
    text in the SAME paragraph, must still be present (byte-identical -
    it's reused verbatim, never rebuilt) after that paragraph is
    translated."""
    source = tmp_path / "source.docx"
    document = docx.Document()
    paragraph = document.add_paragraph()
    _add_symbol_run(paragraph, font="Wingdings", char_code="F0A7")
    paragraph.add_run(" More declassifications will be posted soon.")
    document.save(str(source))

    destination = tmp_path / "result.docx"
    engine = DocxEngine()
    engine.open(str(source))
    translate_document(engine, _FakeProvider(), [], "DE", "en")
    engine.save(str(destination))

    with zipfile.ZipFile(destination) as archive:
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)

    syms = root.findall(f".//{{{_W_NS}}}sym")
    assert len(syms) == 1, "the <w:sym/> bullet did not survive translation"
    assert syms[0].get(f"{{{_W_NS}}}font") == "Wingdings"
    assert syms[0].get(f"{{{_W_NS}}}char") == "F0A7"

    # And the actual .docx is still structurally valid.
    result = docx.Document(str(destination))
    assert "More declassifications" in result.paragraphs[0].text
    assert result.paragraphs[0].text.startswith("[DE]")


def test_bold_toggle_is_inserted_at_the_schema_correct_position(tmp_path: Path) -> None:
    """Regression guard for the rFonts/sz-preservation fix itself: when a
    copied <w:rPr> already carries <w:rFonts>/<w:sz> (in valid CT_RPr
    schema order) and a run's own bold flag adds a fresh <w:b/>, that new
    element must be inserted BEFORE <w:sz> (not appended after it, which
    is invalid per the OOXML schema and has, in manual testing, made Word
    ask to repair a document) - see docx_engine.py's _insert_rpr_child().
    """
    source = tmp_path / "source.docx"
    document = docx.Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Important notice.")
    run.bold = True
    run.font.name = "Montserrat"
    run.font.size = Pt(14)
    document.save(str(source))

    destination = tmp_path / "result.docx"
    engine = DocxEngine()
    engine.open(str(source))
    translate_document(engine, _FakeProvider(), [], "DE", "en")
    engine.save(str(destination))

    with zipfile.ZipFile(destination) as archive:
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)

    # The fake provider's "[DE] " prefix lands outside the <b>...</b> span
    # (see the analogous comment in the hyperlink test above), so the
    # paragraph now has two runs: a leading, non-bold "[DE] " and the
    # actual bold "Important notice." - find that second one specifically.
    bold_run = next(
        r
        for r in root.findall(f".//{{{_W_NS}}}r")
        if "".join(t.text or "" for t in r.findall(f"{{{_W_NS}}}t")).strip() == "Important notice."
    )
    rpr = bold_run.find(f"{{{_W_NS}}}rPr")
    assert rpr is not None

    tags = [etree.QName(child).localname for child in rpr]
    # rFonts, then b/bCs, then sz - never b/bCs appended after sz.
    assert tags.index("rFonts") < tags.index("b") < tags.index("sz")

    # And the result still opens cleanly via python-docx (which parses
    # every part it touches, including validating well-formed XML).
    result = docx.Document(str(destination))
    assert any(run.bold for run in result.paragraphs[0].runs)
