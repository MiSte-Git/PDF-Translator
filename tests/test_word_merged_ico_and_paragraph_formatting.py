"""Covers two 03.09.2026 regressions Michael reported from a real merged,
side-by-side-translated multi-ICO document:

"Jetzt sehe ich das im zweiten zusammengefügten Dokument ist der ober
Teil, der ja im einzelnem Dokument die erste Seite ist, im zusammengefügten
Dokument als normale Seite gehandhabt wird und dadurch doppelt dargestellt
wird. Beim Zusammenfügen nebeneinander müssen wir dadrauf achten, das die
ober Teil der ersten Seite des hinzugefügten Dokuments auch nicht übersetzt
wird. [...] Dann hat der Originale Body Text 11 pt und der Übersetzte 12
pt. Wenn dann sollten beide die gleiche Grösse haben. Dadurch verrutschen
die Absätze."

1. DocxEngine.open(ico_mode=True) used to find the page-1 metadata
   separator shape ONCE for the whole document and treat only paragraphs
   before that single, first occurrence as non-translatable. A merged
   multi-ICO document (pipeline/word/merge.py, one real section break per
   source) has ONE such metadata block PER SOURCE - every source after the
   first had its own metadata block wrongly left translatable=True. Now
   the scan runs per-section (see _section_ranges()) - reproduced here
   with a REAL merge_docx_files() call over the project's own
   representative_ico.docx fixture, not a hand-rolled multi-section
   fixture.
2. A translated run whose formatting came entirely from its paragraph's
   OWN <w:pPr>/<w:rPr> ("paragraph mark run properties" - common in real
   Word documents that set a paragraph's font/size once rather than on
   every individual run) had nothing to base its rebuilt <w:rPr> on and
   silently fell back to the document's raw default size instead - now
   WordParagraph.mark_rpr threads that through as a fallback (see
   pipeline/word/html_bridge.py's ParagraphHtml.base_rpr resolution).
"""
from __future__ import annotations

import zipfile
from pathlib import Path

import docx
import pytest
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from pipeline.translation.base import TranslationResult
from pipeline.word.docx_engine import DocxEngine, _w
from pipeline.word.merge import merge_docx_files
from pipeline.word.translate_document import translate_document

_FIXTURES = Path(__file__).parent / "fixtures"
_REPRESENTATIVE_ICO = _FIXTURES / "representative_ico.docx"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class _FakeProvider:
    def translate_html(self, html: str, target_lang: str, source_lang: str | None = None) -> TranslationResult:
        return TranslationResult(text=f"[{target_lang}] {html}", source_lang=source_lang or "en", target_lang=target_lang, provider="fake")


def test_every_merged_sources_own_page_1_metadata_block_stays_untranslated(tmp_path: Path) -> None:
    merged = tmp_path / "merged.docx"
    merge_docx_files([_REPRESENTATIVE_ICO, _REPRESENTATIVE_ICO, _REPRESENTATIVE_ICO], merged)

    engine = DocxEngine()
    engine.open(str(merged), ico_mode=True)

    assert engine.separator_found is True

    paragraphs = engine.get_paragraphs()
    metadata_texts = [
        "".join(run.text for run in p.runs) for p in paragraphs if not p.translatable
    ]
    # One metadata paragraph per merged source (representative_ico.docx has
    # exactly one, see tests/fixtures/representative_ico.docx) - not just
    # the first source's.
    assert metadata_texts == ["ICO Metadata: Issuer XYZ"] * 3

    # And every OTHER paragraph (the actual body content of all three
    # sources) is still translatable - the fix must not have swung the
    # other way and marked everything non-translatable.
    translatable_count = sum(1 for p in paragraphs if p.translatable)
    assert translatable_count > 0


def test_merged_metadata_paragraphs_are_not_sent_to_the_translation_provider(tmp_path: Path) -> None:
    """The actual end-to-end symptom: translate_document() must skip ALL
    THREE metadata blocks (as "non-translatable"), not just the first."""
    merged = tmp_path / "merged.docx"
    merge_docx_files([_REPRESENTATIVE_ICO, _REPRESENTATIVE_ICO, _REPRESENTATIVE_ICO], merged)

    engine = DocxEngine()
    engine.open(str(merged), ico_mode=True)
    stats = translate_document(engine, _FakeProvider(), [], "DE", "en")

    engine.save(str(tmp_path / "result.docx"), overwrite=True)
    result = docx.Document(str(tmp_path / "result.docx"))
    metadata_paragraphs = [p.text for p in result.paragraphs if p.text == "ICO Metadata: Issuer XYZ"]
    # All three still read exactly "ICO Metadata: Issuer XYZ" - none of
    # them got the fake provider's "[DE] " prefix, proving none were sent
    # to translation.
    assert metadata_paragraphs == ["ICO Metadata: Issuer XYZ"] * 3


def test_a_run_with_no_own_rpr_inherits_size_from_the_paragraph_mark(tmp_path: Path) -> None:
    """03.09.2026 (Michael: "Dann hat der Originale Body Text 11 pt und der
    Übersetzte 12 pt"): a paragraph whose formatting lives entirely in its
    <w:pPr>/<w:rPr> (no individual run carries its own <w:rPr>) must still
    give a rebuilt/translated run something to inherit that size from,
    instead of falling back to the document's raw default.
    """
    source = tmp_path / "source.docx"
    document = docx.Document()
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    mark_rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11pt
    mark_rpr.append(sz)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Montserrat")
    mark_rpr.append(rfonts)
    p_pr.append(mark_rpr)

    # A run with NO <w:rPr> of its own - relies entirely on the paragraph
    # mark's rPr above, exactly like real-world Word documents that set
    # body-text formatting once per paragraph rather than per run.
    run_element = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Body text with no run-level formatting."
    run_element.append(t)
    paragraph._p.append(run_element)
    document.save(str(source))

    destination = tmp_path / "result.docx"
    engine = DocxEngine()
    engine.open(str(source))
    translate_document(engine, _FakeProvider(), [], "DE", "en")
    engine.save(str(destination))

    with zipfile.ZipFile(destination) as archive:
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)

    translated_run = next(
        r
        for r in root.findall(f".//{{{_W_NS}}}r")
        if "Body text" in "".join(t.text or "" for t in r.findall(f"{{{_W_NS}}}t"))
    )
    rpr = translated_run.find(f"{{{_W_NS}}}rPr")
    assert rpr is not None, "translated run has no rPr at all - lost the paragraph-mark-inherited size"

    sz_element = rpr.find(f"{{{_W_NS}}}sz")
    assert sz_element is not None
    assert sz_element.get(f"{{{_W_NS}}}val") == "22"

    rfonts_element = rpr.find(f"{{{_W_NS}}}rFonts")
    assert rfonts_element is not None
    assert rfonts_element.get(f"{{{_W_NS}}}ascii") == "Montserrat"
