"""Covers pipeline/word/side_by_side.py (03.09.2026, Michael: "Können wir
auch noch beim Übersetzen die Option anbieten das der Original und der
Übersetzte Text nebeneinander angezeigt werden?"). Exercises the real
DocxEngine + translate_document() + build_side_by_side_body() pipeline
end-to-end with a fake provider (like tests/test_word_merge.py, .docx
fixtures are built inline via python-docx), then reads the result back
with python-docx to assert on the resulting table structure - not just
that no exception was raised.
"""
from __future__ import annotations

from pathlib import Path

import docx
import pytest
from docx.enum.section import WD_ORIENT

from pipeline.translation.base import TranslationResult
from pipeline.word.docx_engine import DocxEngine
from pipeline.word.side_by_side import build_side_by_side_body, capture_original_paragraph_elements
from pipeline.word.translate_document import translate_document


class _FakeProvider:
    """Deterministic stand-in for a real TranslationProvider - prefixes
    the (plain-text, untagged in these fixtures) HTML so the translated
    text is trivially distinguishable from the original in assertions."""

    def translate_html(self, html: str, target_lang: str, source_lang: str | None = None) -> TranslationResult:
        return TranslationResult(text=f"[{target_lang}] {html}", source_lang=source_lang or "en", target_lang=target_lang, provider="fake")


def _make_docx(path: Path, paragraphs: list[str]) -> Path:
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(str(path))
    return path


def _run_side_by_side(source: Path, destination: Path, target_lang: str = "DE") -> None:
    engine = DocxEngine()
    engine.open(str(source))
    original_elements = capture_original_paragraph_elements(engine)
    translated_runs: dict[int, list] = {}
    translate_document(engine, _FakeProvider(), [], target_lang, "en", translated_runs=translated_runs)
    build_side_by_side_body(engine, original_elements, translated_runs, target_lang)
    engine.save(str(destination))


def test_side_by_side_output_is_a_two_column_table(tmp_path: Path) -> None:
    source = _make_docx(tmp_path / "source.docx", ["First paragraph.", "Second paragraph."])
    destination = tmp_path / "result.docx"

    _run_side_by_side(source, destination)

    result = docx.Document(str(destination))
    assert len(result.tables) == 1
    table = result.tables[0]
    assert len(table.columns) == 2
    # Header row + one row per source paragraph.
    assert len(table.rows) == 1 + 2

    header = table.rows[0]
    assert header.cells[0].text == "Original"
    assert "DE" in header.cells[1].text

    first_row, second_row = table.rows[1], table.rows[2]
    assert first_row.cells[0].text == "First paragraph."
    assert first_row.cells[1].text == "[DE] First paragraph."
    assert second_row.cells[0].text == "Second paragraph."
    assert second_row.cells[1].text == "[DE] Second paragraph."


def test_side_by_side_switches_page_to_landscape(tmp_path: Path) -> None:
    source = _make_docx(tmp_path / "source.docx", ["Just one paragraph."])
    destination = tmp_path / "result.docx"

    original = docx.Document(str(source))
    original_section = original.sections[0]
    assert original_section.orientation == WD_ORIENT.PORTRAIT

    _run_side_by_side(source, destination)

    result = docx.Document(str(destination))
    result_section = result.sections[0]
    assert result_section.orientation == WD_ORIENT.LANDSCAPE
    # Width/height actually swapped, not just the orientation flag.
    assert result_section.page_width == original_section.page_height
    assert result_section.page_height == original_section.page_width


def test_side_by_side_empty_paragraph_becomes_a_spanning_row_not_duplicated(tmp_path: Path) -> None:
    """An empty paragraph is never actually translated (paragraph_to_html()
    returns "" -> translate_document() skips it, see _translate_paragraph())
    - it must show up once, spanning both columns, not as identical empty
    text duplicated into both cells (which would look like a translation
    happened when it didn't).
    """
    source = _make_docx(tmp_path / "source.docx", ["Real text.", ""])
    destination = tmp_path / "result.docx"

    _run_side_by_side(source, destination)

    result = docx.Document(str(destination))
    table = result.tables[0]
    # header + "Real text." (2-column) + "" (1 spanning cell) = 3 rows
    assert len(table.rows) == 3
    spanning_row = table.rows[2]
    # python-docx merges a horizontally-spanned cell's Cell objects into
    # the same underlying <w:tc> - both indices resolve to one cell.
    assert spanning_row.cells[0]._tc is spanning_row.cells[1]._tc


def test_side_by_side_preserves_document_after_zip_roundtrip(tmp_path: Path) -> None:
    """Sanity check the output is a structurally valid .docx (regression
    guard in the same spirit as the merge.py "file corrupt" bug fixed
    03.09.2026): every relationship target referenced from the rebuilt
    body must still exist in the archive, and the archive must open
    cleanly via python-docx (which validates the OPC package on load).
    """
    source = _make_docx(tmp_path / "source.docx", ["Alpha.", "Beta.", "Gamma."])
    destination = tmp_path / "result.docx"

    _run_side_by_side(source, destination)

    # Raises if the .docx is structurally broken - python-docx parses
    # every part it touches (document.xml, styles, sections, tables) on
    # Document() construction and on .tables/.paragraphs access.
    result = docx.Document(str(destination))
    assert len(result.tables) == 1
    assert [p.text for p in result.paragraphs if p.text.strip()] == []


def test_side_by_side_is_opt_in_normal_mode_untouched(tmp_path: Path) -> None:
    """build_side_by_side_body() is never called unless explicitly
    requested - the normal translate_document() pass on its own must
    keep producing the familiar layout-preserving output (a document
    with no table at all)."""
    source = _make_docx(tmp_path / "source.docx", ["Only paragraph."])
    destination = tmp_path / "result.docx"

    engine = DocxEngine()
    engine.open(str(source))
    translate_document(engine, _FakeProvider(), [], "DE", "en")
    engine.save(str(destination))

    result = docx.Document(str(destination))
    assert result.tables == []
    assert result.paragraphs[0].text == "[DE] Only paragraph."
