"""Compare all four translation providers (Google, DeepL, OpenAI, Grok) on
one PDF, side by side, to judge translation quality per provider.

For every translatable paragraph block extracted from the PDF, translates
it with each provider in turn and writes a report with one Original/
Provider section per block - as a Word document (default) or, if
--output points at a .md file (or python-docx isn't installed), as
Markdown. A provider failure (quota, missing credentials, API error, ...)
does not abort the run - the provider's entry just gets "[Nicht
verfuegbar: <message>]" and the script continues with the next
provider/block. Does not modify the PDF itself and does not touch any
provider/engine class - read-only consumer of both.

Usage:
    python tools/compare_providers.py <pdf_pfad> [--template <template.json>] [--output <out.docx|out.md>]
    python tools/compare_providers.py "1526 VIRELICON.pdf"
    python tools/compare_providers.py doc.pdf --template template.json --output vergleich.docx
    python tools/compare_providers.py doc.pdf --output vergleich.md
"""
from __future__ import annotations

import argparse
import html as html_lib
import re
import sys
import time
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from pipeline.pdf.base import TextBlock
from pipeline.pdf.pymupdf_engine import PyMuPdfEngine, spans_to_html
from pipeline.pdf.template import load_json
from pipeline.translation.base import TranslationError, TranslationProvider
from pipeline.translation.deepl_provider import DeepLProvider
from pipeline.translation.google_provider import GoogleTranslateProvider
from pipeline.translation.grok_provider import GrokProvider
from pipeline.translation.openai_provider import OpenAIProvider
from pipeline.translation.protected_terms import (
    derive_protected_term,
    protect_terms,
    restore_terms,
)

# Order here is also the column order in the output table.
PROVIDER_FACTORIES: list[tuple[str, type[TranslationProvider]]] = [
    ("Google", GoogleTranslateProvider),
    ("DeepL", DeepLProvider),
    ("OpenAI", OpenAIProvider),
    ("Grok", GrokProvider),
]

_TAG_RE = re.compile(r"<[^>]+>")


def collect_model_names(providers: dict[str, TranslationProvider]) -> dict[str, str]:
    """Each provider's model_name (e.g. OpenAIProvider.model_name), for
    display in the report header. Not part of the TranslationProvider
    protocol (base.py) - only some providers have a selectable model, the
    others (Google/DeepL) return a fixed API identifier instead.
    """
    return {name: provider.model_name for name, provider in providers.items()}  # type: ignore[attr-defined]


def _is_meaningful_block(text: str) -> bool:
    """False for empty/whitespace-only text or text with no letters at all
    (e.g. a bare page number "12" or a symbol-only line "***").
    """
    stripped = text.strip()
    if not stripped:
        return False
    return any(ch.isalpha() for ch in stripped)


def collect_blocks(engine: PyMuPdfEngine) -> list[TextBlock]:
    """All translatable, non-trivial blocks across every page, in reading
    order. Header/footer/template-zone blocks are already excluded by
    extract_blocks() via translatable=False.
    """
    blocks: list[TextBlock] = []
    for page in engine.get_pages():
        page_blocks = sorted(engine.extract_blocks(page.index), key=lambda block: block.bbox[1])
        for block in page_blocks:
            if block.translatable and _is_meaningful_block(block.text):
                blocks.append(block)
    return blocks


def _html_to_plain(value: str) -> str:
    """Strip HTML tags (from translate_html() output) and collapse
    whitespace to a single line, for display in a Markdown table cell.
    """
    without_tags = _TAG_RE.sub(" ", value)
    return " ".join(html_lib.unescape(without_tags).split())


def translate_block(
    provider: TranslationProvider,
    block: TextBlock,
    target_lang: str,
    source_lang: str,
    protected_terms: list[str],
) -> str:
    """Translate one block with one provider, applying the same
    protected-terms placeholder protection the main pipeline uses. Blocks
    with spans go through translate_html() (formatting-aware, protects
    terms internally); plain blocks (no spans) go through translate() with
    protection applied around the call, since translate() itself takes no
    protected_terms parameter.
    """
    if block.spans:
        html_text = spans_to_html(block.spans)
        result = provider.translate_html(  # type: ignore[attr-defined]
            html_text,
            target_lang=target_lang,
            source_lang=source_lang,
            protected_terms=protected_terms,
        )
        return _html_to_plain(result.text)

    protected_text, mapping = protect_terms(block.text, protected_terms)
    result = provider.translate(protected_text, target_lang=target_lang, source_lang=source_lang)
    return restore_terms(result.text, mapping)


def _format_blockquote(text: str) -> str:
    lines = text.splitlines() or [""]
    return "\n".join(f"> {line}" for line in lines)


def _format_table_cell(text: str) -> str:
    single_line = " ".join(text.split())
    return single_line.replace("|", "\\|")


def write_markdown(
    output_path: Path,
    entries: list[tuple[int, int, str, dict[str, str]]],
    provider_names: list[str],
    pdf_path: Path,
    model_names: dict[str, str],
    run_timestamp: datetime,
) -> None:
    lines: list[str] = [f"# Übersetzungsvergleich: {pdf_path.name}", ""]
    lines.append(f"- Datum: {run_timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
    for name in provider_names:
        lines.append(f"- {name}: {model_names[name]}")
    lines.append("")
    lines.append("---")
    lines.append("")

    for paragraph_num, page_index, original_text, translations in entries:
        lines.append(f"## Absatz {paragraph_num} (Seite {page_index + 1})")
        lines.append("")
        lines.append("**Original:**")
        lines.append(_format_blockquote(original_text))
        lines.append("")
        lines.append("| Provider | Übersetzung |")
        lines.append("|----------|-------------|")
        for name in provider_names:
            lines.append(f"| {name:<8} | {_format_table_cell(translations[name])} |")
        lines.append("")
        lines.append("---")
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")


def _add_horizontal_rule(document) -> None:
    """Add a paragraph with a bottom border - renders as a genuine
    horizontal line in Word, not a run of dash characters. Standard
    python-docx recipe (no direct API for this exists): a "w:pBdr/w:bottom"
    element on an otherwise empty paragraph's properties.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_multiline_run(paragraph, text: str) -> None:
    """Add `text` to `paragraph` as a single run, turning embedded newlines
    into soft line breaks (Run.add_break()) instead of paragraph breaks -
    keeps e.g. a multi-line original quote as one logical paragraph.
    """
    lines = text.splitlines() or [""]
    run = paragraph.add_run(lines[0])
    for line in lines[1:]:
        run.add_break()
        run = paragraph.add_run(line)


def write_docx(
    output_path: Path,
    entries: list[tuple[int, int, str, dict[str, str]]],
    provider_names: list[str],
    pdf_path: Path,
    model_names: dict[str, str],
    run_timestamp: datetime,
) -> None:
    """Write the comparison as a Word document (python-docx). Each block
    (heading, "Original:" label + text, and each provider's label + text)
    gets paragraph_format.keep_with_next=True on every one of its
    paragraphs except the last, so Word avoids splitting a single block
    across a page boundary - without forcing an actual page break, so
    ~100+ blocks don't blow the document up into an unnecessarily long
    one-block-per-page document. Blocks are visually separated by a
    horizontal-rule paragraph instead. Raises ImportError if python-docx
    isn't installed - callers may catch this and fall back to
    write_markdown().
    """
    from docx import Document

    document = Document()

    document.add_heading(f"Übersetzungsvergleich: {pdf_path.name}", level=1)
    date_paragraph = document.add_paragraph()
    date_paragraph.add_run("Datum: ").bold = True
    date_paragraph.add_run(run_timestamp.strftime("%Y-%m-%d %H:%M:%S"))
    for name in provider_names:
        model_paragraph = document.add_paragraph()
        model_paragraph.add_run(f"{name}: ").bold = True
        model_paragraph.add_run(model_names[name])
    _add_horizontal_rule(document)

    for entry_index, (paragraph_num, page_index, original_text, translations) in enumerate(entries):
        block_paragraphs = []

        heading = document.add_heading(f"Absatz {paragraph_num} (Seite {page_index + 1})", level=2)
        block_paragraphs.append(heading)

        original_label = document.add_paragraph()
        original_label.add_run("Original:").bold = True
        block_paragraphs.append(original_label)

        original_paragraph = document.add_paragraph()
        _add_multiline_run(original_paragraph, original_text)
        block_paragraphs.append(original_paragraph)

        for name in provider_names:
            provider_label = document.add_paragraph()
            provider_label.add_run(f"{name}:").bold = True
            block_paragraphs.append(provider_label)

            provider_paragraph = document.add_paragraph()
            _add_multiline_run(provider_paragraph, translations[name])
            block_paragraphs.append(provider_paragraph)

        for paragraph in block_paragraphs[:-1]:
            paragraph.paragraph_format.keep_with_next = True

        if entry_index < len(entries) - 1:
            _add_horizontal_rule(document)

    document.save(str(output_path))


def resolve_output(pdf_path: Path, output_arg: str | None) -> tuple[Path, str]:
    """Decide the output path and format ("docx" or "md"). No --output:
    defaults to "<pdf_stem>_vergleich.docx". --output ending in ".md" ->
    Markdown; any other/no extension -> docx (the default format).
    """
    if output_arg is None:
        return pdf_path.parent / f"{pdf_path.stem}_vergleich.docx", "docx"
    path = Path(output_arg)
    if path.suffix.lower() == ".md":
        return path, "md"
    return path, "docx"


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Vergleicht die Übersetzung eines PDFs durch alle vier Provider (Google, DeepL, OpenAI, Grok)."
    )
    parser.add_argument("pdf_path", help="Pfad zur Quell-PDF")
    parser.add_argument("--template", help="Pfad zu einer JSON-Template-Datei (header_bbox/footer_bbox/first_page_zones)")
    parser.add_argument(
        "--output",
        help="Pfad der Ausgabedatei (Default: <pdf_name>_vergleich.docx). "
        "Endung .md erzeugt Markdown statt Word.",
    )
    parser.add_argument("--source-lang", default="en", help="Quellsprache (Default: en)")
    parser.add_argument("--target-lang", default="de", help="Zielsprache (Default: de)")
    args = parser.parse_args()

    pdf_path = Path(args.pdf_path)
    if not pdf_path.exists():
        print(f"PDF nicht gefunden: {pdf_path}")
        sys.exit(1)

    template = load_json(Path(args.template)) if args.template else None
    output_path, output_format = resolve_output(pdf_path, args.output)

    engine = PyMuPdfEngine(template=template)
    engine.open(str(pdf_path))
    blocks = collect_blocks(engine)

    protected_terms = [derive_protected_term(pdf_path.name)]
    providers: dict[str, TranslationProvider] = {name: factory() for name, factory in PROVIDER_FACTORIES}
    provider_names = list(providers.keys())
    model_names = collect_model_names(providers)
    failure_counts = {name: 0 for name in provider_names}

    total_blocks = len(blocks)
    entries: list[tuple[int, int, str, dict[str, str]]] = []
    run_timestamp = datetime.now()
    start_time = time.monotonic()

    for i, block in enumerate(blocks, start=1):
        translations: dict[str, str] = {}
        for provider_name in provider_names:
            print(f"Block {i}/{total_blocks} – Provider {provider_name}...")
            try:
                translations[provider_name] = translate_block(
                    providers[provider_name], block, args.target_lang, args.source_lang, protected_terms
                )
            except TranslationError as exc:
                failure_counts[provider_name] += 1
                translations[provider_name] = f"[Nicht verfügbar: {exc}]"
        entries.append((i, block.page_index, block.text, translations))

    if output_format == "docx":
        try:
            write_docx(output_path, entries, provider_names, pdf_path, model_names, run_timestamp)
        except ImportError as exc:
            output_path = output_path.with_suffix(".md")
            print(f"python-docx nicht installiert ({exc}) - weiche auf Markdown aus: {output_path}")
            write_markdown(output_path, entries, provider_names, pdf_path, model_names, run_timestamp)
    else:
        write_markdown(output_path, entries, provider_names, pdf_path, model_names, run_timestamp)
    elapsed = time.monotonic() - start_time

    print()
    print("=== Statistik ===")
    print(f"Übersetzte Blöcke: {total_blocks}")
    for name in provider_names:
        print(f"  {name}: {failure_counts[name]} Fehler")
    print(f"Gesamtlaufzeit: {elapsed:.1f}s")
    print(f"Ausgabe: {output_path}")


if __name__ == "__main__":
    main()
